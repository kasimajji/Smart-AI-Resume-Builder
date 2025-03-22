from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
import magic
from docx import Document
from PyPDF2 import PdfReader
import re

app = Flask(__name__)

# Configure upload folder
UPLOAD_FOLDER = 'temp_uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Create upload folder if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_mime_type(file_path):
    mime = magic.Magic(mime=True)
    return mime.from_file(file_path)

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def check_for_tables(file_path, file_type):
    if file_type == 'docx':
        doc = Document(file_path)
        return len(doc.tables) > 0
    # For PDF, we'll use a simple heuristic based on text patterns
    elif file_type == 'pdf':
        text = extract_text_from_pdf(file_path)
        # Look for patterns that might indicate tables
        table_patterns = [
            r'\|.*\|',  # Text between pipe characters
            r'\+[-]+\+',  # ASCII table borders
            r'[\t]{2,}',  # Multiple tabs
            r'[ ]{4,}',   # Multiple spaces
        ]
        for pattern in table_patterns:
            if re.search(pattern, text):
                return True
    return False

def check_for_images(file_path, file_type):
    if file_type == 'docx':
        doc = Document(file_path)
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        return image_count > 0
    # For PDF, check for image objects
    elif file_type == 'pdf':
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                if '/XObject' in page['/Resources']:
                    x_objects = page['/Resources']['/XObject'].get_object()
                    if any(obj['/Subtype'] == '/Image' for obj in x_objects.values()):
                        return True
    return False

def analyze_keywords(text):
    common_resume_keywords = [
        'experience', 'skills', 'education', 'project', 'achievement',
        'leadership', 'management', 'development', 'analysis', 'research',
        'team', 'communication', 'responsibility', 'success', 'improvement'
    ]
    
    found_keywords = []
    text_lower = text.lower()
    
    for keyword in common_resume_keywords:
        if keyword in text_lower:
            found_keywords.append(keyword)
    
    return found_keywords

@app.route('/api/ats/analyze', methods=['POST'])
def analyze_resume():
    if 'file' not in request.files:
        return jsonify({
            'error': 'No file provided'
        }), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({
            'error': 'No file selected'
        }), 400
    
    if not allowed_file(file.filename):
        return jsonify({
            'error': 'Invalid file format. Only PDF and DOCX files are allowed.'
        }), 400
    
    try:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Determine file type
        mime_type = get_mime_type(file_path)
        file_type = 'pdf' if 'pdf' in mime_type.lower() else 'docx'
        
        # Extract text based on file type
        text = extract_text_from_pdf(file_path) if file_type == 'pdf' else extract_text_from_docx(file_path)
        
        # Perform analysis
        has_tables = check_for_tables(file_path, file_type)
        has_images = check_for_images(file_path, file_type)
        keywords_found = analyze_keywords(text)
        
        # Calculate score
        score = 100
        feedback = []
        
        # File format check
        feedback.append({
            'type': 'success',
            'message': f'File format ({file_type.upper()}) is valid and commonly accepted by ATS systems'
        })
        
        # Check for images
        if has_images:
            score -= 15
            feedback.append({
                'type': 'warning',
                'message': 'Images detected: This may reduce ATS compatibility. Consider removing images and using text instead.'
            })
        
        # Check for tables
        if has_tables:
            score -= 20
            feedback.append({
                'type': 'error',
                'message': 'Tables detected: Many ATS systems cannot properly parse tables. Consider using plain text formatting.'
            })
        
        # Keyword analysis
        keyword_score = min(len(keywords_found) * 5, 25)  # Max 25 points for keywords
        score = max(min(score + keyword_score, 100), 0)  # Ensure score stays between 0-100
        
        feedback.append({
            'type': 'info',
            'message': f'Keywords found: {", ".join(keywords_found)}'
        })
        
        # Clean up
        os.remove(file_path)
        
        return jsonify({
            'score': score,
            'feedback': feedback,
            'keywords': keywords_found
        })
        
    except Exception as e:
        # Clean up in case of error
        if os.path.exists(file_path):
            os.remove(file_path)
        
        return jsonify({
            'error': f'Error analyzing resume: {str(e)}'
        }), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)